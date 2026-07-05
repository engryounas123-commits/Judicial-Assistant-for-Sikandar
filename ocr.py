"""
OCR Module – Document Processing Pipeline
Handles PDF, DOCX, images (JPG/PNG) with OCR for Urdu + English
"""

import os
import re
import logging
import tempfile
from pathlib import Path
from typing import Tuple, Dict, List, Optional

import numpy as np
from PIL import Image, ImageFilter, ImageEnhance

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Production-grade document processor supporting:
    - PDF (text-based and scanned)
    - DOCX
    - Images (JPG, PNG, TIFF, BMP)
    - Urdu + English mixed text
    - Page number preservation
    - OCR noise cleaning
    """

    def __init__(self):
        self._init_ocr()
        self._init_cv2()
        self._init_pdf()
        logger.info("DocumentProcessor initialized")

    # ------------------------------------------------------------------
    # Initialization helpers
    # ------------------------------------------------------------------

    def _init_ocr(self):
        try:
            import pytesseract
            self.pytesseract = pytesseract

            # Try common tesseract paths
            for path in [
                "/usr/bin/tesseract",
                "/usr/local/bin/tesseract",
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            ]:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    break

            env_path = os.getenv("TESSERACT_PATH")
            if env_path and os.path.exists(env_path):
                pytesseract.pytesseract.tesseract_cmd = env_path

            self.ocr_available = True
            logger.info("Tesseract OCR ready")
        except ImportError:
            self.pytesseract = None
            self.ocr_available = False
            logger.warning("pytesseract not installed – OCR disabled")

    def _init_cv2(self):
        try:
            import cv2
            self.cv2 = cv2
            self.cv2_available = True
        except ImportError:
            self.cv2 = None
            self.cv2_available = False
            logger.warning("OpenCV not installed – image preprocessing disabled")

    def _init_pdf(self):
        try:
            import pdfplumber
            self.pdfplumber = pdfplumber
            self.pdf_available = True
        except ImportError:
            self.pdfplumber = None
            self.pdf_available = False
            logger.warning("pdfplumber not installed – PDF processing limited")

    # ------------------------------------------------------------------
    # Main public API
    # ------------------------------------------------------------------

    def process_file(self, file_path: str, original_name: str = "") -> Tuple[str, Dict]:
        """
        Process any supported file and return (text, metadata).
        Dispatches to the appropriate handler based on file extension.
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        metadata: Dict = {
            "original_name": original_name or path.name,
            "file_type": suffix,
            "pages": 0,
            "ocr_used": False,
            "word_count": 0,
            "language_detected": "en",
            "error": None,
        }

        try:
            dispatch = {
                ".pdf": self._process_pdf,
                ".docx": self._process_docx,
                ".doc": self._process_docx,
                ".jpg": self._process_image,
                ".jpeg": self._process_image,
                ".png": self._process_image,
                ".tiff": self._process_image,
                ".tif": self._process_image,
                ".bmp": self._process_image,
            }

            handler = dispatch.get(suffix)
            if handler is None:
                metadata["error"] = f"Unsupported file type: {suffix}"
                return f"[Unsupported: {suffix}]", metadata

            text, metadata = handler(file_path, metadata)
            text = self._clean_text(text)
            metadata["word_count"] = len(text.split())
            metadata["language_detected"] = self._detect_language(text)
            logger.info(
                f"Processed '{original_name}': {metadata['word_count']} words, "
                f"lang={metadata['language_detected']}"
            )
            return text, metadata

        except Exception as exc:
            logger.error(f"Error processing '{original_name}': {exc}", exc_info=True)
            metadata["error"] = str(exc)
            return f"[Processing error for {original_name}: {exc}]", metadata

    def merge_documents(self, all_text: Dict) -> str:
        """
        Merge all extracted document texts into one unified case bundle.
        """
        parts: List[str] = []
        parts.append("=" * 70)
        parts.append("   CASE DOCUMENT BUNDLE – AI JUDICIAL BENCH ASSISTANT")
        parts.append("=" * 70)

        for filename, data in all_text.items():
            parts.append(f"\n{'─' * 60}")
            parts.append(f"DOCUMENT : {filename}")
            parts.append(f"TYPE     : {data.get('file_type', 'unknown').upper().lstrip('.')}")
            parts.append(
                f"PAGES    : {data.get('metadata', {}).get('pages', data.get('pages', '?'))}"
            )
            ocr_flag = "YES" if data.get("metadata", {}).get("ocr_used") else "No"
            parts.append(f"OCR      : {ocr_flag}")
            parts.append(f"{'─' * 60}\n")
            parts.append(data.get("text", "[No text extracted]"))

        parts.append(f"\n{'=' * 70}")
        parts.append("   END OF CASE BUNDLE")
        parts.append("=" * 70)
        return "\n".join(parts)

    # ------------------------------------------------------------------
    # Per-format processors
    # ------------------------------------------------------------------

    def _process_pdf(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Handle text-based and scanned PDF files."""
        if not self.pdf_available:
            return "[pdfplumber not installed – cannot process PDF]", metadata

        all_pages: List[str] = []

        try:
            with self.pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = ""

                    # 1) Try native text extraction
                    try:
                        page_text = page.extract_text() or ""
                    except Exception as exc_txt:
                        logger.warning(f"Page {page_num} text extraction failed: {exc_txt}")

                    # 2) Fall back to OCR when text is too sparse
                    if len(page_text.strip()) < 50 and self.ocr_available:
                        logger.info(f"Page {page_num}: sparse text – applying OCR")
                        page_text = self._ocr_pdf_page(page, page_num, page_text)
                        metadata["ocr_used"] = True

                    # 3) Append any tables
                    try:
                        for table in page.extract_tables() or []:
                            if table:
                                rows = [
                                    " | ".join(str(c or "") for c in row)
                                    for row in table
                                    if row
                                ]
                                page_text += "\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]"
                    except Exception:
                        pass

                    all_pages.append(f"[PAGE {page_num}]\n{page_text}")

        except Exception as exc:
            logger.error(f"PDF read error: {exc}", exc_info=True)
            return f"[PDF error: {exc}]", metadata

        return "\n\n".join(all_pages), metadata

    def _ocr_pdf_page(self, page, page_num: int, fallback_text: str) -> str:
        """Convert a pdfplumber page to image and run OCR."""
        try:
            tmp_img = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            tmp_img.close()
            page.to_image(resolution=300).save(tmp_img.name)

            ocr_text, _ = self._process_image(tmp_img.name, {})
            os.unlink(tmp_img.name)

            return ocr_text if len(ocr_text.strip()) > len(fallback_text.strip()) else fallback_text
        except Exception as exc:
            logger.warning(f"OCR fallback failed for page {page_num}: {exc}")
            return fallback_text

    def _process_docx(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Extract text from DOCX documents."""
        try:
            from docx import Document
        except ImportError:
            return "[python-docx not installed]", metadata

        try:
            doc = Document(file_path)
            parts: List[str] = []

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                style_name = para.style.name.lower()
                if "heading 1" in style_name:
                    parts.append(f"\n# {para.text}")
                elif "heading 2" in style_name:
                    parts.append(f"\n## {para.text}")
                elif "heading" in style_name:
                    parts.append(f"\n### {para.text}")
                else:
                    parts.append(para.text)

            for table in doc.tables:
                parts.append("\n[TABLE]")
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
                parts.append("[/TABLE]")

            metadata["pages"] = 1
            return "\n".join(parts), metadata

        except Exception as exc:
            logger.error(f"DOCX error: {exc}", exc_info=True)
            return f"[DOCX error: {exc}]", metadata

    def _process_image(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Run OCR on an image file (JPG/PNG/TIFF/BMP)."""
        if not self.ocr_available:
            return "[Tesseract OCR not installed]", metadata

        try:
            img = Image.open(file_path)
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")

            # Preprocess for better accuracy
            img = self._preprocess_image(img)

            cfg = r"--oem 3 --psm 3"

            # English
            try:
                text_en = self.pytesseract.image_to_string(img, lang="eng", config=cfg)
            except Exception:
                text_en = ""

            # Urdu (optional – only when lang pack is available)
            text_ur = ""
            try:
                text_ur = self.pytesseract.image_to_string(img, lang="urd", config=cfg)
            except Exception:
                pass

            # Merge both language outputs
            if text_ur.strip() and len(text_ur.strip()) > len(text_en.strip()) * 0.3:
                text = (text_en + "\n\n[URDU TEXT]\n" + text_ur).strip()
            else:
                text = text_en

            metadata["ocr_used"] = True
            metadata["pages"] = 1
            return text, metadata

        except Exception as exc:
            logger.error(f"Image OCR error: {exc}", exc_info=True)
            return f"[OCR error: {exc}]", metadata

    # ------------------------------------------------------------------
    # Image preprocessing
    # ------------------------------------------------------------------

    def _preprocess_image(self, img: Image.Image) -> Image.Image:
        """
        Enhance image for better OCR:
        - Convert to grayscale
        - Denoise (via OpenCV or PIL fallback)
        - Adaptive thresholding / binarisation
        - Deskew (small angles only)
        """
        if self.cv2_available:
            return self._cv2_preprocess(img)
        return self._pil_preprocess(img)

    def _cv2_preprocess(self, img: Image.Image) -> Image.Image:
        """OpenCV-based preprocessing pipeline."""
        try:
            arr = np.array(img)
            gray = (
                self.cv2.cvtColor(arr, self.cv2.COLOR_RGB2GRAY)
                if len(arr.shape) == 3
                else arr
            )

            # Denoise
            denoised = self.cv2.fastNlMeansDenoising(gray, h=10)

            # Adaptive threshold
            binary = self.cv2.adaptiveThreshold(
                denoised,
                255,
                self.cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                self.cv2.THRESH_BINARY,
                11,
                2,
            )

            # Deskew (only for small angles to avoid distortion)
            try:
                coords = np.column_stack(np.where(binary < 255))
                if len(coords) > 100:
                    angle = self.cv2.minAreaRect(coords.astype(np.float32))[-1]
                    angle = -(90 + angle) if angle < -45 else -angle
                    if abs(angle) < 10:
                        h, w = binary.shape
                        M = self.cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
                        binary = self.cv2.warpAffine(
                            binary,
                            M,
                            (w, h),
                            flags=self.cv2.INTER_CUBIC,
                            borderMode=self.cv2.BORDER_REPLICATE,
                        )
            except Exception:
                pass  # Deskew failure is non-fatal

            return Image.fromarray(binary)
        except Exception as exc:
            logger.warning(f"cv2 preprocess failed, using PIL fallback: {exc}")
            return self._pil_preprocess(img)

    def _pil_preprocess(self, img: Image.Image) -> Image.Image:
        """PIL-only preprocessing fallback."""
        img = img.convert("L")
        img = ImageEnhance.Contrast(img).enhance(2.0)
        img = img.filter(ImageFilter.SHARPEN)
        return img

    # ------------------------------------------------------------------
    # Text utilities
    # ------------------------------------------------------------------

    def _clean_text(self, text: str) -> str:
        """Normalise and clean extracted text."""
        if not text:
            return ""

        # Collapse excessive blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Collapse multiple spaces
        text = re.sub(r"[ \t]{2,}", " ", text)
        # Remove form-feed characters
        text = text.replace("\f", "\n[PAGE BREAK]\n")
        # Remove most non-printable ASCII while keeping Urdu (U+0600–U+06FF)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        # Remove repeated dashes/pipes (OCR table artifacts)
        text = re.sub(r"-{4,}", "---", text)
        text = re.sub(r"\|{2,}", "|", text)

        return text.strip()

    def _detect_language(self, text: str) -> str:
        """Heuristic language detection (Urdu vs English)."""
        if not text:
            return "en"

        alpha_chars = [c for c in text if c.isalpha()]
        if not alpha_chars:
            return "en"

        urdu_chars = sum(1 for c in alpha_chars if "\u0600" <= c <= "\u06FF")
        ratio = urdu_chars / len(alpha_chars)

        if ratio > 0.5:
            return "ur"
        if ratio > 0.15:
            return "bilingual"
        return "en"
