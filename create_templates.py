"""
Template Creator Script
Run this script once to generate base DOCX templates for each Pakistani court.
Usage: python create_templates.py
"""

import os
from pathlib import Path
from datetime import datetime

TEMPLATES_DIR = Path(__file__).parent / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)


def create_template(filename: str, court_title: str, court_short: str):
    """Create a single court DOCX template."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx")
        return

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # Base style
    normal_style      = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    navy  = RGBColor(0x1A, 0x27, 0x44)
    gold  = RGBColor(0xC6, 0x92, 0x2A)

    # Header line
    hdr_para = doc.sections[0].header.paragraphs[0]
    hdr_para.clear()
    hdr_run = hdr_para.add_run(court_short)
    hdr_run.font.size = Pt(9)
    hdr_run.font.color.rgb = navy
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Court title
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run(court_title)
    r1.bold = True
    r1.font.size = Pt(14)
    r1.font.color.rgb = navy

    doc.add_paragraph()

    # Parties placeholder
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run("[PETITIONER / PLAINTIFF NAME]").bold = True

    vs = doc.add_paragraph()
    vs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vs.add_run("VERSUS").italic = True

    rp = doc.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.add_run("[RESPONDENT / DEFENDANT NAME]").bold = True

    doc.add_paragraph()

    # Case info
    ci = doc.add_paragraph()
    ci.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ci.add_run("Case No. __________ / ____    |    Bench: __________________________")

    doc.add_paragraph()
    hr = doc.add_paragraph("─" * 80)
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Section placeholders
    sections_to_add = [
        "1. INTRODUCTION",
        "2. FACTS OF THE CASE",
        "3. ISSUES FRAMED",
        "4. ARGUMENTS OF THE PARTIES",
        "5. EVIDENCE ANALYSIS",
        "6. LEGAL REASONING",
        "7. PRECEDENTS",
        "8. FINDINGS",
        "9. OPERATIVE ORDER",
    ]
    for sec_title in sections_to_add:
        h = doc.add_paragraph()
        hr_ = h.add_run(sec_title)
        hr_.bold = True
        hr_.font.size = Pt(12)
        hr_.font.color.rgb = navy

        body = doc.add_paragraph()
        body.add_run("[Content to be filled by AI Judgment Writer]").italic = True
        doc.add_paragraph()

    # Footer
    for sec in doc.sections:
        footer = sec.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fr = fp.add_run(
            f"{court_short}  |  AI Judicial Bench Assistant  |  "
            f"Template Version 1.0  |  {datetime.now().year}"
        )
        fr.font.size = Pt(8)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    path = TEMPLATES_DIR / filename
    doc.save(path)
    print(f"✅ Template created: {path}")


def main():
    print("Creating court templates…\n")

    templates = [
        ("supreme.docx",          "IN THE SUPREME COURT OF PAKISTAN",               "Supreme Court of Pakistan"),
        ("lahore_hc.docx",        "IN THE LAHORE HIGH COURT, LAHORE",               "Lahore High Court"),
        ("islamabad_hc.docx",     "IN THE ISLAMABAD HIGH COURT, ISLAMABAD",         "Islamabad High Court"),
        ("peshawar_hc.docx",      "IN THE PESHAWAR HIGH COURT, PESHAWAR",           "Peshawar High Court"),
        ("balochistan_hc.docx",   "IN THE HIGH COURT OF BALOCHISTAN, QUETTA",       "Balochistan High Court"),
        ("sindh_hc.docx",         "IN THE HIGH COURT OF SINDH, KARACHI",            "Sindh High Court"),
        ("fsc.docx",              "IN THE FEDERAL SHARIAT COURT OF PAKISTAN",       "Federal Shariat Court"),
    ]

    for fname, court_title, court_short in templates:
        create_template(fname, court_title, court_short)

    print(f"\n✅ All {len(templates)} templates created in: {TEMPLATES_DIR}")


if __name__ == "__main__":
    main()
