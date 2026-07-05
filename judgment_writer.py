"""
Judgment Writer Module
Generates a fully formatted Microsoft Word (.docx) court judgment
following Pakistani judicial conventions and the selected court's template.
"""

import io
import logging
from datetime import datetime
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)


class JudgmentWriter:
    """
    Converts AI judgment data + case data into a professional DOCX file.
    Returns bytes that can be served as a Streamlit download.
    """

    # Court seals / short names
    COURT_HEADERS: Dict[str, Dict] = {
        "Supreme Court Pakistan": {
            "title": "IN THE SUPREME COURT OF PAKISTAN",
            "bench_label": "PRESENT:",
            "color": "1F3864",
        },
        "Lahore High Court": {
            "title": "IN THE LAHORE HIGH COURT, LAHORE",
            "bench_label": "BENCH:",
            "color": "1F3864",
        },
        "Islamabad High Court": {
            "title": "IN THE ISLAMABAD HIGH COURT, ISLAMABAD",
            "bench_label": "BENCH:",
            "color": "1F3864",
        },
        "Peshawar High Court": {
            "title": "IN THE PESHAWAR HIGH COURT, PESHAWAR",
            "bench_label": "BENCH:",
            "color": "1F3864",
        },
        "Balochistan High Court": {
            "title": "IN THE HIGH COURT OF BALOCHISTAN, QUETTA",
            "bench_label": "BENCH:",
            "color": "1F3864",
        },
        "Sindh High Court": {
            "title": "IN THE HIGH COURT OF SINDH, KARACHI",
            "bench_label": "BENCH:",
            "color": "1F3864",
        },
        "Federal Shariat Court": {
            "title": "IN THE FEDERAL SHARIAT COURT OF PAKISTAN, ISLAMABAD",
            "bench_label": "BENCH:",
            "color": "1F3864",
        },
    }

    def generate_docx(
        self,
        case_data: Dict,
        judgment: Dict,
        opinions: Dict[str, str],
        precedents: List[Dict],
        court: str,
        judgment_style: str,
        language: str = "English",
    ) -> bytes:
        """
        Build and return the complete judgment DOCX as bytes.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        doc = Document()
        self._setup_page(doc, Cm, Inches)
        self._apply_base_styles(doc, Pt, RGBColor)

        court_info = self.COURT_HEADERS.get(court, self.COURT_HEADERS["Supreme Court Pakistan"])
        parties = case_data.get("parties", {})
        identifiers = case_data.get("case_identifiers", {})

        # ---- Page 1: Dashboard / Coversheet ----
        self._add_dashboard_page(doc, case_data, judgment, court, court_info, parties, identifiers, Pt, RGBColor, WD_ALIGN_PARAGRAPH)

        doc.add_page_break()

        # ---- Judgment Body ----
        self._add_judgment_header(doc, court_info, parties, identifiers, judgment, judgment_style, Pt, RGBColor, WD_ALIGN_PARAGRAPH)

        self._add_section(doc, "1. INTRODUCTION", self._build_introduction(case_data, judgment, court), Pt, RGBColor)
        self._add_section(doc, "2. FACTS OF THE CASE", self._build_facts(case_data, judgment), Pt, RGBColor)
        self._add_section(doc, "3. ISSUES FRAMED", self._build_issues(judgment), Pt, RGBColor)
        self._add_section(doc, "4. ARGUMENTS OF THE PARTIES", self._build_arguments(case_data, judgment), Pt, RGBColor)
        self._add_section(doc, "5. EVIDENCE ANALYSIS", self._build_evidence(case_data, judgment), Pt, RGBColor)
        self._add_section(doc, "6. LEGAL REASONING", self._build_legal_reasoning(judgment), Pt, RGBColor)
        self._add_section(doc, "7. PRECEDENTS CONSIDERED", self._build_precedents(judgment, precedents), Pt, RGBColor)
        self._add_section(doc, "8. FINDINGS", self._build_findings(judgment), Pt, RGBColor)

        # Shariah section (only if relevant)
        shariah = judgment.get("shariah_aspects", "")
        if shariah and shariah not in ("N/A", "Not applicable", ""):
            self._add_section(doc, "9. SHARIAH / ISLAMIC LAW CONSIDERATIONS", shariah, Pt, RGBColor)
            self._add_operative_order(doc, judgment, parties, Pt, RGBColor, WD_ALIGN_PARAGRAPH, section_num=10)
        else:
            self._add_operative_order(doc, judgment, parties, Pt, RGBColor, WD_ALIGN_PARAGRAPH, section_num=9)

        # Appendix: Judge Panel Opinions
        if opinions:
            doc.add_page_break()
            self._add_appendix_opinions(doc, opinions, Pt, RGBColor)

        # Footer / disclaimer
        self._add_footer(doc, court, Pt)

        # Save to bytes
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        logger.info(f"DOCX judgment generated ({buf.getbuffer().nbytes} bytes)")
        return buf.read()

    # ------------------------------------------------------------------
    # Page setup & styles
    # ------------------------------------------------------------------

    def _setup_page(self, doc, Cm, Inches):
        from docx.shared import Cm
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    def _apply_base_styles(self, doc, Pt, RGBColor):
        """Apply base Normal style."""
        normal = doc.styles["Normal"]
        font   = normal.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        pf = normal.paragraph_format
        pf.space_after  = Pt(6)
        pf.line_spacing = Pt(18)

    # ------------------------------------------------------------------
    # Dashboard / Coversheet
    # ------------------------------------------------------------------

    def _add_dashboard_page(self, doc, case_data, judgment, court, court_info, parties, identifiers, Pt, RGBColor, WD_ALIGN):
        from docx.shared import RGBColor as RGB, Pt as DPt

        def _gold():    return RGB(0xC6, 0x92, 0x2A)
        def _navy():    return RGB(0x1a, 0x27, 0x44)
        def _white():   return RGB(0xFF, 0xFF, 0xFF)

        # Court title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN.CENTER
        run = p.add_run("⚖  AI JUDICIAL BENCH ASSISTANT  ⚖")
        run.bold = True
        run.font.size = DPt(16)
        run.font.color.rgb = _navy()

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN.CENTER
        r2 = p2.add_run(court_info["title"])
        r2.bold = True
        r2.font.size = DPt(13)
        r2.font.color.rgb = _navy()

        doc.add_paragraph()

        # Case snapshot table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        def add_row(label, value):
            row = table.add_row()
            lc, vc = row.cells
            lp = lc.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = DPt(11)
            lr.font.color.rgb = _navy()
            vp = vc.paragraphs[0]
            vr = vp.add_run(str(value or "N/A"))
            vr.font.size = DPt(11)

        add_row("Petitioner / Plaintiff", parties.get("plaintiff_petitioner", "N/A"))
        add_row("Respondent / Defendant", parties.get("defendant_respondent", "N/A"))
        add_row("Court", court)
        add_row("Case Number", identifiers.get("case_number", "N/A"))
        add_row("Year", identifiers.get("year", str(datetime.now().year)))
        add_row("Case Type", str(judgment.get("case_type", "N/A")).title())
        add_row("Decision", judgment.get("decision", "RESERVED"))
        add_row("Confidence Score", f"{judgment.get('confidence_score', 0)}%")
        add_row("Generated On", datetime.now().strftime("%d %B %Y, %H:%M"))

        doc.add_paragraph()

        # Key Issues
        issues = judgment.get("issues_framed", [])
        if issues:
            ih = doc.add_paragraph()
            ihr = ih.add_run("KEY ISSUES FRAMED")
            ihr.bold = True
            ihr.font.size = DPt(12)
            ihr.font.color.rgb = _navy()

            for i, issue in enumerate(issues[:5], 1):
                p_issue = doc.add_paragraph(style="List Number")
                p_issue.add_run(issue).font.size = DPt(11)

        # Operative Order summary box
        op = judgment.get("operative_order", "")
        if op:
            doc.add_paragraph()
            ophead = doc.add_paragraph()
            ohr = ophead.add_run("OPERATIVE ORDER (Summary)")
            ohr.bold = True
            ohr.font.color.rgb = _navy()
            opdoc = doc.add_paragraph()
            opr = opdoc.add_run(op[:600] + ("…" if len(op) > 600 else ""))
            opr.font.size = DPt(11)
            opr.italic = True

        # Disclaimer
        doc.add_paragraph()
        disc = doc.add_paragraph()
        disc.alignment = WD_ALIGN.CENTER
        dr = disc.add_run(
            "AI-GENERATED DRAFT | FOR REFERENCE ONLY | NOT A BINDING COURT ORDER"
        )
        dr.bold = True
        dr.font.size = DPt(9)
        dr.font.color.rgb = RGB(0x99, 0x33, 0x00)

    # ------------------------------------------------------------------
    # Judgment header (after page break)
    # ------------------------------------------------------------------

    def _add_judgment_header(self, doc, court_info, parties, identifiers, judgment, style, Pt, RGBColor, WD_ALIGN):
        from docx.shared import RGBColor as RGB, Pt as DPt

        navy = RGB(0x1a, 0x27, 0x44)

        # Court name
        ct = doc.add_paragraph()
        ct.alignment = WD_ALIGN.CENTER
        ctr = ct.add_run(court_info["title"])
        ctr.bold = True
        ctr.font.size = DPt(14)
        ctr.font.color.rgb = navy

        doc.add_paragraph()

        # Parties
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN.CENTER
        ppr = pp.add_run(parties.get("plaintiff_petitioner", "Petitioner/Plaintiff").upper())
        ppr.bold = True
        ppr.font.size = DPt(12)

        vs = doc.add_paragraph()
        vs.alignment = WD_ALIGN.CENTER
        vsr = vs.add_run("VERSUS")
        vsr.italic = True
        vsr.font.size = DPt(11)

        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN.CENTER
        rpr = rp.add_run(parties.get("defendant_respondent", "Respondent/Defendant").upper())
        rpr.bold = True
        rpr.font.size = DPt(12)

        doc.add_paragraph()

        # Case meta
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN.CENTER
        case_no = identifiers.get("case_number", "N/A")
        year    = identifiers.get("year", str(datetime.now().year))
        mr = meta.add_run(f"Case No. {case_no} / {year}  |  Style: {style}")
        mr.font.size = DPt(11)
        mr.font.color.rgb = RGB(0x44, 0x44, 0x44)

        # AI Bench
        bench = doc.add_paragraph()
        bench.alignment = WD_ALIGN.CENTER
        br = bench.add_run(
            "CORAM: AI Judge Panel (Law Expert, Shariah Judge, Domain Specialist, "
            "Precedent Researcher, Chief Justice AI)"
        )
        br.font.size = DPt(10)
        br.italic = True

        # Date
        dt = doc.add_paragraph()
        dt.alignment = WD_ALIGN.CENTER
        dtr = dt.add_run(f"Date of Judgment: {datetime.now().strftime('%d %B %Y')}")
        dtr.font.size = DPt(11)

        doc.add_paragraph()

        # Horizontal rule substitute
        hr = doc.add_paragraph("─" * 80)
        hr.alignment = WD_ALIGN.CENTER
        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Generic section adder
    # ------------------------------------------------------------------

    def _add_section(self, doc, heading: str, content: str, Pt, RGBColor):
        from docx.shared import RGBColor as RGB, Pt as DPt

        h = doc.add_paragraph()
        hr = h.add_run(heading)
        hr.bold = True
        hr.font.size = DPt(12)
        hr.font.color.rgb = RGB(0x1a, 0x27, 0x44)

        if content:
            for para_text in content.split("\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = DPt(18)
                p.add_run(para_text).font.size = DPt(11)

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Content builders
    # ------------------------------------------------------------------

    def _build_introduction(self, case_data: Dict, judgment: Dict, court: str) -> str:
        parties = case_data.get("parties", {})
        pp = parties.get("plaintiff_petitioner", "the Petitioner")
        rd = parties.get("defendant_respondent", "the Respondent")
        case_type = judgment.get("case_type", "general").title()
        summary = judgment.get("case_summary", "")

        return (
            f"This {case_type} matter came before {court} on the petition/suit filed by "
            f"{pp} against {rd}. "
            f"{summary}\n\n"
            f"The AI Judicial Bench, comprising five specialist judge agents, deliberated "
            f"upon the facts, evidence, applicable Pakistani law, Shariah principles, and "
            f"binding precedents before arriving at this judgment."
        )

    def _build_facts(self, case_data: Dict, judgment: Dict) -> str:
        lines = []
        key_facts = judgment.get("key_facts") or case_data.get("key_facts", [])
        if key_facts:
            lines.append("The following facts emerge from the case record:")
            for i, fact in enumerate(key_facts, 1):
                lines.append(f"{i}. {fact}")

        timeline = case_data.get("timeline", [])
        if timeline:
            lines.append("\nCHRONOLOGY:")
            for ev in timeline[:8]:
                lines.append(f"• {ev.get('date','')} – {ev.get('event','')}")

        return "\n".join(lines) if lines else "Facts to be determined from the case record."

    def _build_issues(self, judgment: Dict) -> str:
        issues = judgment.get("issues_framed", [])
        if not issues:
            return "Issues to be framed based on the pleadings."
        return "\n".join(f"{i}. {iss}" for i, iss in enumerate(issues, 1))

    def _build_arguments(self, case_data: Dict, judgment: Dict) -> str:
        pp = case_data.get("parties", {}).get("plaintiff_petitioner", "Petitioner")
        rd = case_data.get("parties", {}).get("defendant_respondent", "Respondent")
        relief = case_data.get("relief_sought", "")

        text = (
            f"A. ARGUMENTS OF {pp.upper()} (Petitioner/Plaintiff)\n\n"
            f"The learned counsel for the {pp} argued that their client is entitled to "
            f"relief on the grounds set forth in the petition/plaint. "
        )
        if relief:
            text += f"Specifically, relief was sought as follows: {relief[:400]}\n\n"

        text += (
            f"B. ARGUMENTS OF {rd.upper()} (Respondent/Defendant)\n\n"
            f"The learned counsel for the {rd} opposed the petition/suit and contended "
            f"that the claims of the petitioner/plaintiff are without merit and prayed "
            f"for dismissal of the case."
        )
        return text

    def _build_evidence(self, case_data: Dict, judgment: Dict) -> str:
        ev_analysis = judgment.get("evidence_analysis", "")
        evidence_items = case_data.get("evidence_items", [])
        witnesses = case_data.get("witnesses", [])

        lines = []
        if ev_analysis:
            lines.append(ev_analysis)
        if evidence_items:
            lines.append("\nEVIDENCE ON RECORD:")
            for item in evidence_items[:10]:
                lines.append(f"• {item}")
        if witnesses:
            lines.append("\nWITNESSES:")
            for w in witnesses:
                lines.append(f"• {w.get('id','')} – {w.get('name','')}")

        return "\n".join(lines) if lines else "Evidence to be assessed from the case record."

    def _build_legal_reasoning(self, judgment: Dict) -> str:
        reasoning = judgment.get("legal_reasoning", "")
        laws = judgment.get("applicable_laws", [])

        text = reasoning or "Legal reasoning to be completed."
        if laws:
            text += "\n\nAPPLICABLE LAWS:\n"
            text += "\n".join(f"• {law}" for law in laws)
        return text

    def _build_precedents(self, judgment: Dict, precedents: List[Dict]) -> str:
        lines = []
        applied = judgment.get("precedents_applied", [])

        if applied:
            lines.append("The following precedents were considered and applied:")
            for p in applied:
                ver = "VERIFIED" if p.get("verified") else "REQUIRES VERIFICATION"
                lines.append(
                    f"\n{p.get('citation','N/A')} [{ver}]\n"
                    f"Principle: {p.get('principle','')}"
                )

        rag_preds = [r for r in precedents if r.get("type") == "precedent"]
        if rag_preds:
            lines.append("\nADDITIONAL PRECEDENTS RETRIEVED:")
            for p in rag_preds[:4]:
                lines.append(
                    f"• {p.get('title','')} – {p.get('citation','')} "
                    f"[{p.get('court','')}]: {p.get('legal_principle','')[:150]}"
                )

        return "\n".join(lines) if lines else "No specific precedents cited."

    def _build_findings(self, judgment: Dict) -> str:
        findings = judgment.get("findings", [])
        if not findings:
            return "Findings to be determined."
        return "\n".join(f"{i}. {f}" for i, f in enumerate(findings, 1))

    # ------------------------------------------------------------------
    # Operative Order
    # ------------------------------------------------------------------

    def _add_operative_order(self, doc, judgment: Dict, parties: Dict, Pt, RGBColor, WD_ALIGN, section_num: int = 9):
        from docx.shared import RGBColor as RGB, Pt as DPt

        navy  = RGB(0x1a, 0x27, 0x44)
        gold  = RGB(0xC6, 0x92, 0x2A)

        # Section heading
        h = doc.add_paragraph()
        hr = h.add_run(f"{section_num}. OPERATIVE ORDER")
        hr.bold = True
        hr.font.size = DPt(12)
        hr.font.color.rgb = navy

        # Decision
        decision = judgment.get("decision", "RESERVED")
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN.CENTER
        dr = dp.add_run(f"ORDER: {decision}")
        dr.bold = True
        dr.font.size = DPt(14)
        dr.font.color.rgb = gold

        doc.add_paragraph()

        # Formal order text
        op_text = judgment.get("operative_order", "")
        if op_text:
            for para in op_text.split("\n"):
                para = para.strip()
                if para:
                    op = doc.add_paragraph()
                    op.paragraph_format.first_line_indent = DPt(18)
                    op.add_run(para).font.size = DPt(11)

        # Relief granted
        relief = judgment.get("relief_granted", [])
        if relief:
            rh = doc.add_paragraph()
            rhr = rh.add_run("RELIEF GRANTED:")
            rhr.bold = True
            rhr.font.size = DPt(11)
            for r in relief:
                rp = doc.add_paragraph(style="List Bullet")
                rp.add_run(r).font.size = DPt(11)

        # Directives
        directives = judgment.get("directives", [])
        if directives:
            dh = doc.add_paragraph()
            dhr = dh.add_run("COURT DIRECTIVES:")
            dhr.bold = True
            dhr.font.size = DPt(11)
            for d in directives:
                dp2 = doc.add_paragraph(style="List Bullet")
                dp2.add_run(d).font.size = DPt(11)

        # Costs
        costs = judgment.get("costs", "")
        if costs:
            cp = doc.add_paragraph()
            cp.add_run(f"COSTS: {costs}").font.size = DPt(11)

        # Signature block
        doc.add_paragraph()
        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN.RIGHT
        sigr = sig.add_run(
            f"Chief Justice AI\n{datetime.now().strftime('%d %B %Y')}\n{doc.sections[0].header.paragraphs[0].text if doc.sections[0].header.paragraphs else 'AI Judicial Bench'}"
        )
        sigr.font.size = DPt(11)
        sigr.italic = True

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Appendix: Judge Panel Opinions
    # ------------------------------------------------------------------

    def _add_appendix_opinions(self, doc, opinions: Dict[str, str], Pt, RGBColor):
        from docx.shared import RGBColor as RGB, Pt as DPt

        navy = RGB(0x1a, 0x27, 0x44)

        ah = doc.add_paragraph()
        ahr = ah.add_run("APPENDIX – JUDGE PANEL DELIBERATIONS")
        ahr.bold = True
        ahr.font.size = DPt(13)
        ahr.font.color.rgb = navy

        ap = doc.add_paragraph()
        ap.add_run(
            "The following individual opinions were expressed by the AI judge panel "
            "during deliberation and are appended for reference:"
        ).font.size = DPt(11)

        for agent_name, opinion in opinions.items():
            doc.add_paragraph()
            oh = doc.add_paragraph()
            ohr = oh.add_run(f"[ {agent_name.upper()} ]")
            ohr.bold = True
            ohr.font.size = DPt(11)
            ohr.font.color.rgb = navy

            for line in opinion.split("\n")[:40]:
                line = line.strip()
                if line:
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = DPt(18)
                    p.add_run(line).font.size = DPt(10)

    # ------------------------------------------------------------------
    # Footer
    # ------------------------------------------------------------------

    def _add_footer(self, doc, court: str, Pt):
        from docx.shared import Pt as DPt

        for section in doc.sections:
            footer = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.clear()
            fr = fp.add_run(
                f"AI Judicial Bench Assistant  |  {court}  |  "
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
                "For Research Purposes Only"
            )
            fr.font.size = DPt(8)
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
